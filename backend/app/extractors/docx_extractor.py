from pathlib import Path
from typing import List
import docx

from app.extractors.base import BaseExtractor
from app.schemas.extraction import (
    BlockType,
    DocumentType,
    ExtractedBlock,
    ExtractedImageInfo,
    ExtractionResult,
)
from app.services.image_service import image_service
from app.utils.text_cleaner import clean_extracted_text


class DOCXExtractor(BaseExtractor):
    async def extract(self, file_path: Path, session_id: str) -> ExtractionResult:
        doc = docx.Document(file_path)
        blocks: List[ExtractedBlock] = []
        extracted_images: List[ExtractedImageInfo] = []

        image_service.clear_session_registry(session_id)

        # 1. Extract Images
        image_part_map = {}
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_part = rel.target_part
                image_bytes = image_part.blob
                ext = image_part.content_type.split("/")[-1]

                image_info = image_service.save_image_bytes(
                    session_id=session_id,
                    image_bytes=image_bytes,
                    ext_or_format=ext
                )
                image_part_map[rel.target_ref] = image_info
                
                if not image_info.is_duplicate:
                    extracted_images.append(image_info)

        # 2. Iterate Body Elements
        for element in doc.element.body:
            if element.tag.endswith("p"):
                p = docx.text.paragraph.Paragraph(element, doc)
                text = clean_extracted_text(p.text)

                has_image = False
                for rel_ref, img_info in image_part_map.items():
                    if rel_ref in element.xml:
                        blocks.append(
                            ExtractedBlock(
                                block_type=BlockType.IMAGE,
                                text=None,
                                image_info=img_info
                            )
                        )
                        has_image = True

                if not text and has_image:
                    continue

                if not text:
                    continue

                style_name = p.style.name if p.style else ""

                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1
                    blocks.append(
                        ExtractedBlock(
                            block_type=BlockType.HEADING,
                            text=text,
                            level=level,
                            metadata={"style": style_name},
                        )
                    )
                elif "List" in style_name or text.startswith(("- ", "* ", "• ")):
                    blocks.append(
                        ExtractedBlock(
                            block_type=BlockType.LIST_ITEM,
                            text=text,
                            metadata={"style": style_name},
                        )
                    )
                elif "Code" in style_name:
                    blocks.append(
                        ExtractedBlock(
                            block_type=BlockType.CODE,
                            text=text,
                            metadata={"style": style_name},
                        )
                    )
                else:
                    blocks.append(
                        ExtractedBlock(
                            block_type=BlockType.PARAGRAPH,
                            text=text,
                            metadata={"style": style_name},
                        )
                    )

            elif element.tag.endswith("tbl"):
                t = docx.table.Table(element, doc)
                matrix = []
                for row in t.rows:
                    row_data = [clean_extracted_text(cell.text) for cell in row.cells]
                    matrix.append(row_data)

                if matrix:
                    blocks.append(
                        ExtractedBlock(
                            block_type=BlockType.TABLE,
                            text=None,
                            metadata={"rows": matrix},
                        )
                    )

        return ExtractionResult(
            session_id=session_id,
            document_type=DocumentType.DOCX,
            title=file_path.stem,
            blocks=blocks,
            images=extracted_images,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "unique_image_count": len(extracted_images)
            },
        )